"""docx → 段落级 Chunk；过长段落再切。"""
from __future__ import annotations
from app.core.types import Chunk, Modality, SourceType
from app.core.registry import register
from app.core.exceptions import PreprocessError
from app.preprocess.base import BasePreprocessor
from app.preprocess.text_loader import split_text
from app.utils.hashing import sha256_file, make_chunk_id
from app.utils.io import ensure_absolute
from config import settings


@register((".docx",))
class DocxLoader(BasePreprocessor):
    SUPPORTED_EXT = (".docx",)

    def load(self, file_path: str) -> list[Chunk]:
        path = ensure_absolute(file_path)
        try:
            from docx import Document  # python-docx
        except ImportError as e:
            raise PreprocessError("缺少 python-docx 依赖") from e

        try:
            doc = Document(path)
        except Exception as e:
            raise PreprocessError(f"解析 docx 失败: {path}: {e}") from e

        file_hash = sha256_file(path)
        max_len = settings.SPLITTER.split_length
        chunks: list[Chunk] = []
        idx = 0

        # 段落
        for p_idx, para in enumerate(doc.paragraphs, start=1):
            t = (para.text or "").strip()
            if not t:
                continue
            if len(t) <= max_len:
                chunks.append(Chunk(
                    chunk_id=make_chunk_id(path, idx),
                    source_path=path,
                    source_type=SourceType.DOCX,
                    modality=Modality.TEXT,
                    file_hash=file_hash,
                    page=p_idx,            # paragraph_index 作为 page
                    text=t,
                    preview=t[:200],
                    extra={"from": "paragraph"},
                ))
                idx += 1
            else:
                for ss, ee, seg in split_text(t):
                    chunks.append(Chunk(
                        chunk_id=make_chunk_id(path, idx),
                        source_path=path,
                        source_type=SourceType.DOCX,
                        modality=Modality.TEXT,
                        file_hash=file_hash,
                        page=p_idx,
                        text=seg,
                        preview=seg[:200],
                        extra={"from": "paragraph", "char_start": ss, "char_end": ee},
                    ))
                    idx += 1

        # 表格
        for t_idx, table in enumerate(doc.tables, start=1):
            for r_idx, row in enumerate(table.rows, start=1):
                cells = [c.text.strip() for c in row.cells]
                joined = " | ".join(c for c in cells if c)
                if not joined:
                    continue
                chunks.append(Chunk(
                    chunk_id=make_chunk_id(path, idx),
                    source_path=path,
                    source_type=SourceType.DOCX,
                    modality=Modality.TEXT,
                    file_hash=file_hash,
                    text=joined,
                    preview=joined[:200],
                    extra={"from": "table", "table_idx": t_idx, "row": r_idx},
                ))
                idx += 1

        return chunks
